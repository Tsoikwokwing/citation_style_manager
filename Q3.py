# CMT114 Coursework
# student number:C1937813

import docx
import os, sys
from docx import Document


def change_style(filepath, style='IEEE'):
    # YOUR CODE HERE
    # YOUR CODE HERE
    # YOUR CODE HERE
    if style=='IEEE':
        with open('APAexample_no_hyperlinks.docx','rb') as f:
            doc = Document(f)

        #find out all the In-Text citations, and append them into list_for_dict with numerical order
        In_text_dict = {}
        value = 1
        list_for_dict = []
        checkedlist = []
        for i in range(len(doc.paragraphs)):
            #extract all the In-Text citation except from the References list
            if doc.paragraphs[i].text == 'References':
                break
            else:
                #find out how many left brackets in each paragraph
                content = doc.paragraphs[i].text
                left_bracket_amount = 0
                beg = 0
                while content.find('(', beg) != -1:
                    left_bracket_amount += 1
                    beg = content.find('(', beg) + 1

                #extrac all the extract all the In-Text citation
                start_index = -1
                for n in range(left_bracket_amount):
                    start_index = content.find('(', start_index+1)
                    comma_index = content.find(',', start_index)
                    surname = content[start_index+1:comma_index]
                    year = content[start_index+1:start_index+5]


                    if year.isdigit():  #e.g. surname (YYYY)
                        if content[start_index-5:start_index-1] == ' al.':  #e.g. surname et al. (YYYY)
                            i = 6
                            while i > 5:
                                if (content[start_index - i] == ' ') or (content[start_index - i].islower()):
                                    i += 1
                                else:
                                    break
                            surname = content[start_index-i:start_index-1]
                            key = '(' + surname + ', ' + year + ')'
                            if key not in checkedlist:
                                list_for_dict.append([key, '[' + str(value) + ']'])
                                value += 1
                                checkedlist.append(key)
                        else:
                            i = 2
                            while i > 1:
                                if (content[start_index - i] == ' ') or (content[start_index - i].islower()):
                                    i += 1
                                else:
                                    break
                            if content[start_index-(i+4):start_index-i] == 'and ':  #e.g. surname and surname (YYYY)
                                n = i + 1
                                while n > i:
                                    if (content[start_index - n] == ' ') or (content[start_index - n].islower()):
                                        n += 1
                                    else:
                                        break
                                surname = content[start_index-n:start_index-1].strip()
                                key = '(' + surname + ', ' + year + ')'
                                if key not in checkedlist:
                                    list_for_dict.append([key, '[' + str(value) + ']'])
                                    value += 1
                                    checkedlist.append(key)
                            else:  #e.g. surname (YYYY)
                                surname = content[start_index-i:start_index-1]
                                key = '(' + surname + ', ' + year + ')'
                                if key not in checkedlist:
                                    list_for_dict.append([key, '[' + str(value) + ']'])
                                    value += 1
                                    checkedlist.append(key)

                    elif year == 'e.g.':
                        end_index = content.find(')', start_index) + 1
                        key = content[start_index:end_index].replace('e.g., ','')
                        if key not in checkedlist:
                            if key.find(';') == -1:
                                list_for_dict.append([key, '[' + str(value) + ']'])
                                value += 1
                                checkedlist.append(key)
                            elif key.count(';') == 1:
                                key1 = key.split(';')[0] + ')'
                                if key1 not in checkedlist:
                                    #continue
                                #else:
                                    list_for_dict.append([key1, '[' + str(value) + ']'])
                                    value += 1
                                    checkedlist.append(key1)
                                key2 = '(' + key.split(';')[1].strip()
                                if key2 not in checkedlist:
                                #    continue
                                #else:
                                    list_for_dict.append([key2, '[' + str(value) + ']'])
                                    value += 1
                                    checkedlist.append(key2)
                            elif key.count(';') == 2:
                                key1 = key.split(';')[0] + ')'
                                if key1 not in checkedlist:
                                    #continue
                                #else:
                                    list_for_dict.append([key1, '[' + str(value) + ']'])
                                    value += 1
                                    checkedlist.append(key1)
                                key2 = '(' + key.split(';')[1].strip() + ')'
                                if key2 not in checkedlist:
                                    #continue
                                #else:
                                    list_for_dict.append([key2, '[' + str(value) + ']'])
                                    value += 1
                                    checkedlist.append(key2)
                                key3 = '(' + key.split(';')[2].strip()
                                if key3 not in checkedlist:
                                #    continue
                                #else:
                                    list_for_dict.append([key3, '[' + str(value) + ']'])
                                    value += 1
                                    checkedlist.append(key3)

                    else: #year.isdigit() == False:  #e.g. (surname, YYYY)
                        if surname.find(' ') == -1:  #for In_Text citation like '(Willingham, 2008)'
                            if surname.istitle():
                                end_index = content.find(')', comma_index) + 1
                                key = content[start_index:end_index]

                                if key not in checkedlist:
                                #    continue
                                #else:
                                    if key.find(';') == -1:
                                        list_for_dict.append([key, '[' + str(value) + ']'])
                                        value += 1
                                        checkedlist.append(key)
                                    else:  #e.g.(Moore, 1996; Bushong, 2002)
                                        key1 = key.split(';')[0] + ')'
                                        if key1 not in checkedlist:
                                            #continue
                                        #else:
                                            list_for_dict.append([key1, '[' + str(value) + ']'])
                                            value += 1
                                            checkedlist.append(key1)
                                        key2 = '(' + key.split(';')[1].strip()
                                        if key2 not in checkedlist:
                                        #    continue
                                        #else:
                                            list_for_dict.append([key2, '[' + str(value) + ']'])
                                            value += 1
                                            checkedlist.append(key2)

                        else:  #for In_Text citation like '(Kahneman et al., 1982)' or '(Moynihan and Soderlind, 2003)'
                            space_index = content.find(' ', start_index)
                            surname = content[start_index+1:space_index]
                            if surname.istitle():
                                end_index = content.find(')', comma_index) + 1
                                key = content[start_index:end_index]

                                if key not in checkedlist:
                                #    continue
                                #else:
                                    if key.find(';') == -1:
                                        list_for_dict.append([key, '[' + str(value) + ']'])
                                        value += 1
                                        checkedlist.append(key)
                                    else:  #e.g.(Martin et al., 1993; Weinstein, 2000)
                                        key1 = key.split(';')[0] + ')'
                                        if key1 not in checkedlist:
                                        #    continue
                                        #else:
                                            list_for_dict.append([key1, '[' + str(value) + ']'])
                                            value += 1
                                            checkedlist.append(key1)
                                        key2 = '(' + key.split(';')[1].strip()
                                        if key2 not in checkedlist:
                                        #    continue
                                        #else:
                                            list_for_dict.append([key2, '[' + str(value) + ']'])
                                            value += 1
                                            checkedlist.append(key2)

        #create a dictionary from list_for_dict
        In_text_dict = dict(list_for_dict)

        #replace APA In-Text citations in to IEEE ones part1
        for i in range(len(doc.paragraphs)):
            newcontent = doc.paragraphs[i].text
            '''if newcontent == 'References':
                break
            else:'''
            if '(e.g., ' in newcontent:
                egsta_index = -1
                for x in range(newcontent.count('(e.g., ')):
                    egsta_index = newcontent.find('(e.g., ', egsta_index + 1)
                    egend_index = newcontent.find(')', egsta_index)
                    citation = newcontent[egsta_index:egend_index+1]
                    newcontent = newcontent.replace(citation, citation[:7] + '(' + citation[7:] + ')')

                if newcontent.find(';') != -1:
                    semicolon_index = -1
                    for y in range(newcontent.count(';')):
                        semicolon_index = newcontent.find(';', semicolon_index + 1)
                        if newcontent[semicolon_index+2].isupper(): #replace '; ' into '), (' except ';however'
                            frontpart = newcontent[:semicolon_index]
                            rearpart = newcontent[semicolon_index+2:]
                            newcontent = frontpart + '), (' + rearpart

                for n in range(len(checkedlist)):
                    if checkedlist[n] in newcontent:
                        newcontent = newcontent.replace(checkedlist[n], In_text_dict[checkedlist[n]])
                    else:
                        continue
                    doc.paragraphs[i].clear()
                    doc.paragraphs[i].add_run(newcontent).font.name = 'Times New Roman'

            else:
                if newcontent.find(';') != -1:
                    semicolon_index = -1
                    for y in range(newcontent.count(';')):
                        semicolon_index = newcontent.find(';', semicolon_index + 1)
                        if newcontent[semicolon_index+2].isupper(): #replace '; ' into '), (' except ';however'
                            frontpart = newcontent[:semicolon_index]
                            rearpart = newcontent[semicolon_index+2:]
                            newcontent = frontpart + '), (' + rearpart

                for n in range(len(checkedlist)):
                    if checkedlist[n] in newcontent:
                        newcontent = newcontent.replace(checkedlist[n], In_text_dict[checkedlist[n]])
                    else:
                        continue
                    doc.paragraphs[i].clear()
                    doc.paragraphs[i].add_run(newcontent).font.name = 'Times New Roman'



        #replace APA In-Text citations in to IEEE ones part2
        for i in range(len(doc.paragraphs)):
            newcontent = doc.paragraphs[i].text
            if newcontent == 'References':
                break
            else:
                left_bracket_amount = 0
                beg = 0
                while newcontent.find('(', beg) != -1:
                    left_bracket_amount += 1
                    beg = newcontent.find('(', beg) + 1

                start_index = -1
                for n in range(left_bracket_amount):
                    start_index = newcontent.find('(', start_index+1)
                    #comma_index = content.find(',', start_index)
                    #surname = content[start_index+1:comma_index]
                    year = newcontent[start_index+1:start_index+5]
                    if year.isdigit():  #e.g. surname (YYYY)
                        if newcontent[start_index-5:start_index-1] == ' al.':  #e.g. surname et al. (YYYY)
                            z = 6
                            while z > 5:
                                if (newcontent[start_index - z] == ' ') or (newcontent[start_index - z].islower()):
                                    z += 1
                                else:
                                    break
                            surname = newcontent[start_index-z:start_index-1]
                            newcontent = newcontent.replace(surname + ' (' + year + ')', '(' + surname + ', ' + year + ')')
                        #    doc.paragraphs[i].clear()
                        #    doc.paragraphs[i].add_run(newcontent).font.name = 'Times New Roman'
                        else:
                            z = 2
                            while z > 1:
                                if (newcontent[start_index - z] == ' ') or (newcontent[start_index - z].islower()):
                                    z += 1
                                else:
                                    break
                            if newcontent[start_index-(z+4):start_index-z] == 'and ':  #e.g. surname and surname (YYYY)
                                q = z + 1
                                while q > z:
                                    if (newcontent[start_index - q] == ' ') or (newcontent[start_index - q].islower()):
                                        q += 1
                                    else:
                                        break
                                surname = newcontent[start_index-q:start_index-1].strip()
                                newcontent = newcontent.replace(surname + ' (' + year + ')', '(' + surname + ', ' + year + ')')
                        #        doc.paragraphs[i].clear()
                        #        doc.paragraphs[i].add_run(newcontent).font.name = 'Times New Roman'
                            else:  #e.g. surname (YYYY)
                                surname = newcontent[start_index-z:start_index-1]
                                newcontent = newcontent.replace(surname + ' (' + year + ')', '(' + surname + ', ' + year + ')')
                    #doc.paragraphs[i].clear()
                    #doc.paragraphs[i].add_run(newcontent).font.name = 'Times New Roman'

                for n in range(len(checkedlist)):
                    if checkedlist[n] in newcontent:
                        newcontent = newcontent.replace(checkedlist[n], In_text_dict[checkedlist[n]])
                    else:
                        continue
                    doc.paragraphs[i].clear()
                    doc.paragraphs[i].add_run(newcontent).font.name = 'Times New Roman'


        #convert References list from APA to IEEE
        i = -1
        while doc.paragraphs[i].text != 'References':
            newcontent = doc.paragraphs[i].text
            comma_index = newcontent.find(',')
            end_index = newcontent.find(')')
            surname = newcontent[:comma_index]
            year = newcontent[end_index-4:end_index]
            tempstring = '(' + surname + ', ' + year + ')'
            for n in range(len(checkedlist)):
                if year == checkedlist[n][-5:-1]:
                    if tempstring == checkedlist[n]:
                        order = In_text_dict[checkedlist[n]]
                        if len(order) == 3:
                            newcontent = order[1] + '. ' + newcontent
                        elif len(order) == 4:
                            newcontent = order[1:3] + '. ' + newcontent
                    elif surname == checkedlist[n][1:len(surname)+1]:
                        order = In_text_dict[checkedlist[n]]
                        if len(order) == 3:
                            newcontent = order[1] + '. ' + newcontent
                        elif len(order) == 4:
                            newcontent = order[1:3] + '. ' + newcontent
                    doc.paragraphs[i].clear()
                    doc.paragraphs[i].add_run(newcontent).font.name = 'Times New Roman'
            i -= 1


        #create a list (APA_references_origin) to store the origin references list
        APA_references_origin = []
        idx = 0
        while doc.paragraphs[idx].text != 'References':
            idx += 1
        end_idx = len(doc.paragraphs) - 1

        for i in range(idx+1, end_idx):
            #if doc.paragraphs[i].text != '':
            APA_references_origin.append(doc.paragraphs[i].text)

        #create a new list (APA_references_new) with numerical order
        APA_references_new = []
        for num in range(1, len(checkedlist)+1):
            for i in range(len(APA_references_origin)):
                dot_index = APA_references_origin[i].find('.')
                if APA_references_origin[i][:dot_index] == str(num) :
                    APA_references_new.append(APA_references_origin[i])


        #replace the origin references list into new one (IEEE format)
        n = 0
        for i in range(idx+1, end_idx):
            doc.paragraphs[i].clear()
            doc.paragraphs[i].add_run(APA_references_new[n]).font.name = 'Times New Roman'
            n += 1


        #put this code in the last
        #doc.save('APAexample_no_hyperlinks_test.docx')
        doc.save('APAexample_no_hyperlinks_IEEE_style.docx')

    elif style=='APA':
        with open('IEEEexample_no_hyperlinks.docx','rb') as f:
            doc = Document(f)

        #Convert IEEE to APA
        IEEE_references_origin = []

        #extract all the references into a list:
        i = -1
        while doc.paragraphs[i].text != 'References':
            IEEE_references_origin.append(doc.paragraphs[i].text)
            i -= 1

        IEEE_references_new = IEEE_references_origin[:]
        IEEE_references_new.reverse()

        #delete the number
        for i in range(len(IEEE_references_new)):
            if len(IEEE_references_new[i]) > 0:
                IEEE_references_new[i] = IEEE_references_new[i][3:-1]

        templist = IEEE_references_new[0:-1]
        templist.sort()
        templist.append(IEEE_references_new[-1])
        IEEE_references_new = templist[:]

        #put APA-format references into doc
        i = -1
        while doc.paragraphs[i].text != 'References':
            doc.paragraphs[i].clear()
            doc.paragraphs[i].add_run(IEEE_references_new[i])
            i -= 1


        #generate a apa dictionary
        res = []
        for i in range(1,10):
            if len(IEEE_references_origin[i]) > 0:
                tempstring = IEEE_references_origin[i]
                num = tempstring[0]
                year = tempstring[-5:-1]
                if int(num) in [1,2,3,6,7]:
                    endindex1 = tempstring.find(',')
                    startindex1 = tempstring[:endindex1].rfind('.') + 1
                    surname1 = tempstring[startindex1:endindex1].strip()

                    templist1 = []
                    templist1.append(num)
                    templist1.append('(' + surname1 + ', ' + year + ')')
                elif int(num) in [4,5,8]:
                    endindex1 = tempstring.find(',')
                    startindex1 = tempstring[:endindex1].rfind('.') + 1
                    surname1 = tempstring[startindex1:endindex1].strip()

                    endindex2 = tempstring.find(',', endindex1 + 1)
                    startindex2 = tempstring[:endindex2].rfind('.') + 1
                    surname2 = tempstring[startindex2:endindex2].strip()

                    templist1 = []
                    templist1.append(num)
                    templist1.append('(' + surname1 + ' & ' + surname2 + ', ' + year + ')')
                elif int(num) in [9]:
                    endindex1 = tempstring.find(',')
                    startindex1 = tempstring[:endindex1].rfind('.') + 1
                    surname1 = tempstring[startindex1:endindex1].strip()

                    endindex2 = tempstring.find(',', endindex1 + 1)
                    startindex2 = tempstring[:endindex2].rfind('.') + 1
                    surname2 = tempstring[startindex2:endindex2].strip()

                    endindex3 = tempstring.find(',', endindex2 + 1)
                    startindex3 = tempstring[:endindex3].rfind('.') + 1
                    surname3 = tempstring[startindex3:endindex3].strip()

                    endindex4 = tempstring.find(',', endindex3 + 1)
                    startindex4 = tempstring[:endindex4].rfind('.') + 1
                    surname4 = tempstring[startindex4:endindex4].strip()

                    endindex5 = tempstring.find(',', endindex4 + 1)
                    startindex5 = tempstring[:endindex5].rfind('.') + 1
                    surname5 = tempstring[startindex5:endindex5].strip()

                    templist1 = []
                    templist1.append(num)
                    templist1.append('(' + surname1 + ', ' + surname2 + ', ' + surname3 + ', ' + surname4 + ', & ' + surname5 + ', ' + year +')')

                res.append(templist1)

        apa_dict = dict(res)


        for i in range(len(doc.paragraphs)):
            newcontent = doc.paragraphs[i].text
            for n in range(1,10):
                if '[' + str(n) +']' in newcontent:
                    newcontent = newcontent.replace('[' + str(n) +']', apa_dict[str(n)])
                else:
                    continue
                doc.paragraphs[i].clear()
                doc.paragraphs[i].add_run(newcontent)

        #put this code in the last
        doc.save('IEEEexample_no_hyperlinks_APA_style.docx')

# ---- DO NOT CHANGE THE CODE BELOW ----
if __name__ == "__main__":
    if len(sys.argv)<3: raise ValueError('Provide filename and style as input arguments')
    filepath, style = sys.argv[1], sys.argv[2]
    print('filepath is "{}"'.format(filepath))
    print('target style is "{}"'.format(style))
    change_style(filepath, style)
